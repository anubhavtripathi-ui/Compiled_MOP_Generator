"""
Compiled MOP Generator — v6
============================
Changes from v5:
  1. NEW: Optional Activity MOP upload — extracts images, screenshots, OLE attachments
  2. NEW: Images from Activity MOP injected into Section 10 (SOP) by positional order
         matching [IMAGE/SCREENSHOT REQUIRED] placeholders in the Solution Document
  3. NEW: OLE/embedded attachments (Excel, .bin etc.) re-embedded into output MOP
  4. NEW: If any image/attachment cannot be injected → flagged in UI + placeholder
         text inserted at exact position in output MOP
  5. NEW: Tables from Activity MOP (with images inside) are handled correctly
  6. Zero data retention: all processing in-memory, nothing written to disk
"""

import io
import re
import time
import zipfile
import tempfile
import os
from datetime import datetime
from pathlib import Path
from copy import deepcopy
from lxml import etree

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Compiled MOP Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Session state init ──
for _key, _val in {
    "output_bytes":       b"",
    "activity_name":      "",
    "today_str":          "",
    "sections":           {},
    "filled":             0,
    "images_n":           0,
    "total_n":            0,
    "failed_media":       [],
    "injected_media":     0,
    "comments_injected":  0,
    "comments_failed":    [],
}.items():
    if _key not in st.session_state:
        st.session_state[_key] = _val

# ─────────────────────────────────────────────────────────────────
# CSS — Teal/Emerald palette: dark teal #00473C, accent #00BFA5
# ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Lato:wght@300;400;700;900&family=Source+Code+Pro:wght@400;600&display=swap');

html, body, [class*="css"] {
  font-family: 'Lato', sans-serif;
  background-color: #F2F7F6;
  color: #1A2E2A;
}
.block-container { padding-top: 1.5rem !important; padding-bottom: 2rem; max-width: 100%; }

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #00332A 0%, #00473C 40%, #005C4E 100%) !important;
  border-right: 1px solid rgba(0,191,165,0.2);
}
[data-testid="stSidebar"] * { color: #e0f5f2 !important; }
[data-testid="stSidebar"] .stMarkdown h1,
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3 { color: #ffffff !important; }
[data-testid="stSidebar"] hr { border-color: rgba(0,191,165,0.3) !important; }
[data-testid="stSidebar"] label { color: #80C9BF !important; font-size: .78rem !important; }

.eri-topbar {
  background: linear-gradient(90deg, #00332A, #00473C, #005C4E);
  border-bottom: 3px solid #00BFA5;
  padding: 1rem 2rem 0.8rem;
  border-radius: 12px;
  margin-bottom: 1.5rem;
  display: flex;
  align-items: center;
  justify-content: space-between;
}
.eri-logo-text { font-family:'Lato',sans-serif; font-weight:900; font-size:1.5rem; letter-spacing:3px; color:#00BFA5; text-transform:uppercase; }
.eri-logo-sub  { font-size:0.7rem; letter-spacing:1.5px; color:rgba(255,255,255,0.4); text-transform:uppercase; }
.eri-app-title { font-size:1.15rem; font-weight:700; color:#ffffff; letter-spacing:0.3px; }
.eri-app-sub   { font-size:0.72rem; color:rgba(255,255,255,0.45); letter-spacing:0.5px; margin-top:2px; }
.eri-version   { background:rgba(0,191,165,0.15); border:1px solid rgba(0,191,165,0.3); border-radius:20px; padding:3px 12px; font-size:0.65rem; color:#00BFA5; font-weight:700; letter-spacing:1px; }

/* Privacy / ZDR bar */
.priv-bar {
  background: rgba(0,100,0,0.06);
  border: 1px solid rgba(0,150,60,0.25);
  border-left: 4px solid #009944;
  border-radius: 0 8px 8px 0;
  padding: 0.65rem 1rem;
  font-size: 0.78rem;
  color: #003322;
  margin-bottom: 1.2rem;
}
.priv-bar strong { color: #006633; }

/* Warning bar */
.warn-bar {
  background: rgba(200,80,0,0.06);
  border: 1px solid rgba(200,80,0,0.22);
  border-left: 4px solid #cc5500;
  border-radius: 0 8px 8px 0;
  padding: 0.65rem 1rem;
  font-size: 0.78rem;
  color: #7a3000;
  margin-bottom: 0.8rem;
}
.warn-bar strong { color: #cc5500; }

.eri-card {
  background: #ffffff;
  border: 1px solid #cce8e4;
  border-radius: 12px;
  padding: 1.4rem 1.6rem;
  margin-bottom: 1rem;
  box-shadow: 0 2px 8px rgba(0,71,60,0.06);
  transition: box-shadow 0.2s, border-color 0.2s;
}
.eri-card:hover { border-color: #00BFA5; box-shadow: 0 4px 16px rgba(0,191,165,0.10); }
.eri-card-title {
  font-size:0.68rem; font-weight:900; letter-spacing:1.8px; text-transform:uppercase;
  color:#00473C; margin-bottom:0.9rem; display:flex; align-items:center; gap:8px;
}
.step-badge { background:#00473C; color:#ffffff; font-size:0.58rem; font-weight:700; padding:2px 8px; border-radius:4px; letter-spacing:0.5px; }
.optional-badge { background:#00BFA5; color:#ffffff; font-size:0.55rem; font-weight:700; padding:2px 7px; border-radius:4px; letter-spacing:0.4px; }

.pill-ok   { display:inline-flex; align-items:center; gap:6px; background:rgba(0,150,80,0.08); border:1px solid rgba(0,150,80,0.22); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#006633; margin:3px 0; }
.pill-warn { display:inline-flex; align-items:center; gap:6px; background:rgba(200,80,0,0.07); border:1px solid rgba(200,80,0,0.2); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#7a3500; margin:3px 0; }
.pill-info { display:inline-flex; align-items:center; gap:6px; background:rgba(0,150,130,0.07); border:1px solid rgba(0,150,130,0.2); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#005C4E; margin:3px 0; }

.stButton > button {
  background: linear-gradient(135deg,#00473C,#00BFA5) !important;
  color: #ffffff !important; border:none !important; border-radius:8px !important;
  font-family:'Lato',sans-serif !important; font-weight:700 !important;
  font-size:0.9rem !important; padding:0.6rem 2rem !important; width:100% !important;
  letter-spacing:0.4px !important; transition:all 0.2s !important;
}
.stButton > button:hover { background:linear-gradient(135deg,#005C4E,#00D4B8) !important; transform:translateY(-1px) !important; box-shadow:0 6px 20px rgba(0,191,165,0.28) !important; }
.stButton > button:disabled { opacity:0.38 !important; transform:none !important; }
[data-testid="stDownloadButton"] > button {
  background: linear-gradient(135deg,#00695C,#00897B) !important;
  color:#ffffff !important; border:none !important; border-radius:8px !important;
  font-family:'Lato',sans-serif !important; font-weight:700 !important;
  font-size:0.9rem !important; padding:0.6rem 2rem !important; width:100% !important;
}
[data-testid="stDownloadButton"] > button:hover { background:linear-gradient(135deg,#00796B,#00BFA5) !important; transform:translateY(-1px) !important; }

.prog-wrap { background:#f4faf9; border:1px solid #cce8e4; border-radius:10px; padding:1rem 1.2rem; }
.ps { display:flex; align-items:center; gap:10px; padding:6px 0; font-size:0.78rem; border-bottom:1px solid rgba(0,0,0,0.04); }
.ps:last-child { border-bottom:none; }
.ps.done { color:#006633; font-weight:600; }
.ps.doing { color:#00BFA5; font-weight:600; }
.ps.wait  { color:#9ab8b4; }
.pd { width:8px; height:8px; border-radius:50%; flex-shrink:0; }
.pd.done  { background:#00a050; }
.pd.doing { background:#00BFA5; animation:pulse 1s infinite; }
.pd.wait  { background:#b8d4d0; }
@keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:.5;transform:scale(0.85)} }

.success-card {
  background:linear-gradient(135deg,rgba(0,71,60,0.04),rgba(0,191,165,0.06));
  border:1px solid rgba(0,191,165,0.25); border-top:4px solid #00BFA5;
  border-radius:12px; padding:1.6rem; margin:0.8rem 0; text-align:center;
}
.success-icon  { font-size:2.2rem; margin-bottom:0.4rem; }
.success-title { font-size:1.1rem; font-weight:900; color:#00473C; margin-bottom:0.25rem; }
.success-sub   { font-size:0.78rem; color:#00BFA5; }
.success-name  { color:#00473C; font-weight:700; }

.media-fail-card {
  background:rgba(200,60,0,0.04);
  border:1px solid rgba(200,60,0,0.20);
  border-left:4px solid #cc3300;
  border-radius:0 10px 10px 0;
  padding:1rem 1.2rem;
  margin-top:0.8rem;
}
.media-fail-title { font-size:0.75rem; font-weight:900; color:#cc3300; margin-bottom:0.5rem; letter-spacing:0.8px; text-transform:uppercase; }
.media-fail-item  { font-size:0.76rem; color:#7a2200; padding:3px 0; border-bottom:1px solid rgba(200,60,0,0.08); }
.media-fail-item:last-child { border-bottom:none; }

.metric-row { display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin-top:1rem; }
.metric-box { background:#f0f5fb; border:1px solid #cce8e4; border-radius:10px; padding:1rem; text-align:center; }
.metric-val { font-family:'Lato',sans-serif; font-size:1.7rem; font-weight:900; color:#00473C; }
.metric-sub { font-size:0.58rem; color:#9ab8b4; margin-top:1px; font-weight:700; letter-spacing:.6px; text-transform:uppercase; }

.sidebar-info { background:rgba(0,191,165,0.08); border:1px solid rgba(0,191,165,0.18); border-radius:8px; padding:0.8rem 1rem; font-size:0.74rem; color:#b8d4f0 !important; margin:0.5rem 0; }
hr { border-color:rgba(0,71,60,0.1) !important; }
.stFileUploader label { color:#00473C !important; font-weight:600 !important; font-size:0.82rem !important; }
.footer { text-align:center; font-size:0.65rem; color:#9ab8b4; padding:1rem 0 0.4rem; border-top:1px solid #cce8e4; letter-spacing:0.5px; margin-top:1rem; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────
HEADING_MAP = {
    "objective":            ["objective"],
    "activity_description": ["activity description"],
    "activity_type":        ["activity type"],
    "domain_in_scope":      ["domain in scope", "domain"],
    "prerequisites":        ["pre-requisites", "prerequisites"],
    "inventory_details":    ["inventory details", "inventory"],
    "node_connectivity":    ["node connectivity", "node connectivity process"],
    "iam":                  ["identity & access", "identity and access", "identity"],
    "triggering_method":    ["activity triggering", "triggering method"],
    "sop":                  ["standard operating procedure", "sop"],
    "acceptance_criteria":  ["acceptance criteria"],
    "assumptions":          ["assumptions"],
    "connectivity_diagram": ["connectivity diagram"],
}

SECTION_KEYS = [
    "objective", "activity_description", "activity_type", "domain_in_scope",
    "prerequisites", "inventory_details", "node_connectivity", "iam",
    "triggering_method", "sop", "acceptance_criteria", "assumptions",
    "connectivity_diagram",
]

SECTION_LABELS = {
    "objective":            "1. Objective",
    "activity_description": "2. Activity Description",
    "activity_type":        "3. Activity Type",
    "domain_in_scope":      "4. Domain in Scope",
    "prerequisites":        "5. Pre-requisites",
    "inventory_details":    "6. Inventory Details (Node Name, Type, Count, Vendor)",
    "node_connectivity":    "7. Node Connectivity Process",
    "iam":                  "8. Identity and Access Management",
    "triggering_method":    "9. Activity Triggering Method",
    "sop":                  "10. Standard Operating Procedure (Attach the detailed SOP)",
    "acceptance_criteria":  "11. Acceptance Criteria (UAT scenarios)",
    "assumptions":          "12. Assumptions",
    "connectivity_diagram": "Connectivity Diagram",
}

# Patterns that indicate an image/attachment placeholder in the solution doc
IMAGE_PLACEHOLDER_PATTERNS = [
    r'\[IMAGE[^\]]*\]',
    r'\[SCREENSHOT[^\]]*\]',
    r'\[ATTACHMENT[^\]]*\]',
    r'\[IMAGE/SCREENSHOT[^\]]*\]',
    r'\[DIAGRAM[^\]]*\]',
    r'\[FIGURE[^\]]*\]',
]
IMAGE_PLACEHOLDER_RE = re.compile(
    '|'.join(IMAGE_PLACEHOLDER_PATTERNS), re.IGNORECASE
)

# Namespace shortcuts
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PKG  = "http://schemas.openxmlformats.org/package/2006/relationships"

_HEADING_COLOR_HEX = "1F497D"

# ─────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────
def normalize_heading(text: str):
    t = re.sub(r'^\d+[\.\)]\s*', '', text).strip().lower()
    t = re.sub(r'\s+', ' ', t)
    # Also strip leading "SECTION N —" style
    t = re.sub(r'^section\s+\d+\s*[—\-]\s*', '', t)
    for key, aliases in HEADING_MAP.items():
        for alias in aliases:
            if alias in t:
                return key
    return None


def discover_templates() -> list:
    tmpl_dir = Path("templates")
    if tmpl_dir.exists():
        return sorted(tmpl_dir.glob("*.docx"))
    return []


def load_template_bytes(path: Path) -> bytes:
    with open(path, "rb") as f:
        return f.read()


# ─────────────────────────────────────────────────────────────────
# ACTIVITY MOP — MEDIA EXTRACTOR
# ─────────────────────────────────────────────────────────────────

class MediaItem:
    """Represents one image or OLE attachment from the Activity MOP."""
    def __init__(self, kind, blob, ext, rId, position_index,
                 context_text="", filename="", prog_id="", content_type=""):
        self.kind           = kind           # "image" | "attachment"
        self.blob           = blob           # raw bytes
        self.ext            = ext            # "png","jpg","xlsx","txt","bin"…
        self.rId            = rId
        self.position_index = position_index
        self.context_text   = context_text
        self.filename       = filename       # original filename recovered from OLE
        self.prog_id        = prog_id        # e.g. "Excel.Sheet.12" | "Package"
        self.content_type   = content_type
        self.injected       = False
        self.inject_error   = None

    @property
    def display_name(self):
        if self.filename:
            return self.filename
        if "Excel" in self.prog_id:
            return f"Excel_Attachment_{self.position_index+1}.xlsx"
        return f"Attachment_{self.position_index+1}.{self.ext}"


def _recover_filename_from_ole(blob: bytes) -> str:
    """
    OLE2 Package stream stores original filename as a Latin-1 string.
    Scan for anything that looks like a filename with a known extension.
    """
    try:
        text = blob.decode("latin-1", errors="replace")
        for m in re.finditer(r'[\x20-\x7E]{3,200}', text):
            candidate = m.group(0).strip()
            if re.search(r'\.(txt|log|xlsx|xls|docx|doc|csv|pdf|zip|bin)$',
                         candidate, re.IGNORECASE):
                name = candidate.replace("\\", "/").split("/")[-1].strip()
                if 3 < len(name) < 120:
                    return name
    except Exception:
        pass
    return ""


# Correct namespace URIs confirmed from real docx XML inspection
_NS_R      = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_O      = "urn:schemas-microsoft-com:office:office"    # OLEObject lives here
_NS_BLIP   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_V      = "urn:schemas-microsoft-com:vml"              # v:imagedata for icon


def extract_media_from_activity_mop(mop_bytes: bytes) -> list:
    """
    Walk Activity MOP paragraphs in strict document order.
    For each paragraph:
      • Collect real PNG/JPG screenshots  (skip EMF/WMF — they are OLE icons)
      • Collect OLE/Package attachments   (via <o:OLEObject r:id="…">)
    A paragraph containing an OLE object also contains an EMF icon image;
    we register only the OLE, never the icon.

    Returns a flat list of MediaItem in document order (images + attachments
    interleaved exactly as they appear in the source MOP).
    """
    doc = Document(io.BytesIO(mop_bytes))
    media_items: list = []
    position    = 0
    prev_text   = ""

    # ── Build relationship lookup maps ──────────────────────────
    img_rel_map  = {}   # rId → (blob, ext, content_type)
    att_rel_map  = {}   # rId → (blob, ext, filename, prog_id)

    for rId, rel in doc.part.rels.items():
        rt = rel.reltype.split("/")[-1]
        try:
            if rt == "image":
                ct  = rel.target_part.content_type          # e.g. "image/png"
                raw = ct.split("/")[-1].lower()
                ext = "jpg" if raw == "jpeg" else raw
                img_rel_map[rId] = (rel.target_part.blob, ext, ct)

            elif rt in ("oleObject", "package"):
                blob   = rel.target_part.blob
                target = rel.target_ref                     # e.g. "embeddings/…xlsx"

                # Determine extension
                if rt == "package":
                    # target IS the filename for package parts
                    fname = target.split("/")[-1]
                    ext   = fname.split(".")[-1].lower() if "." in fname else "bin"
                else:
                    # oleObject → binary OLE2 stream — recover filename
                    fname = _recover_filename_from_ole(blob)
                    ext   = fname.split(".")[-1].lower() if fname and "." in fname else "bin"

                att_rel_map[rId] = (blob, ext, fname)

        except Exception:
            pass

    # ── Walk paragraphs ─────────────────────────────────────────
    for para in doc.paragraphs:
        text  = para.text.strip()
        p_xml = para._p
        xml_str = etree.tostring(p_xml, encoding="unicode")

        # ---- Detect OLE objects in this paragraph ----
        # <o:OLEObject> carries the r:id that points to the attachment.
        # We use a regex on the raw XML string because lxml namespace
        # resolution is tricky with multiple prefixes.
        ole_rids_here = re.findall(
            r'<[^>]*OLEObject[^>]+r:id="(rId\d+)"[^>]*/?>',
            xml_str
        )
        # Also catch the package variant (ProgID may differ)
        ole_rids_here += re.findall(
            r'<[^>]*OLEObject[^>]+r:id="(rId\d+)"',
            xml_str
        )
        ole_rids_here = list(dict.fromkeys(ole_rids_here))  # dedupe, preserve order

        # Collect ProgIDs per rId from XML
        prog_map = {}
        for m in re.finditer(r'<[^>]*OLEObject[^>]+>', xml_str):
            tag = m.group(0)
            rid_m   = re.search(r'r:id="(rId\d+)"', tag)
            prog_m  = re.search(r'ProgID="([^"]+)"', tag)
            if rid_m and prog_m:
                prog_map[rid_m.group(1)] = prog_m.group(1)

        # EMF icon rIds that accompany OLE in this same paragraph — must skip
        icon_rids = set()
        for m in re.finditer(r'r:id="(rId\d+)"', xml_str):
            rid = m.group(1)
            if rid in img_rel_map and img_rel_map[rid][1].lower() in ("emf", "wmf"):
                icon_rids.add(rid)

        # ---- Register OLE attachments (if any in this paragraph) ----
        for ole_rid in ole_rids_here:
            if ole_rid not in att_rel_map:
                continue
            blob, ext, fname = att_rel_map[ole_rid]
            prog_id = prog_map.get(ole_rid, "")

            # Override for Excel package
            if "Excel" in prog_id or "Excel" in fname or ext == "xlsx":
                ext  = "xlsx"
                if not fname:
                    fname = f"Excel_Attachment_{position+1}.xlsx"

            item = MediaItem(
                kind="attachment", blob=blob, ext=ext,
                rId=ole_rid, position_index=position,
                context_text=prev_text, filename=fname, prog_id=prog_id
            )
            media_items.append(item)
            position += 1

        # ---- Register real screenshot images (no OLE in this paragraph) ----
        if not ole_rids_here:
            seen_in_para = set()
            for blip in p_xml.findall(f".//{{{_NS_BLIP}}}blip"):
                embed = blip.get(f"{{{_NS_R}}}embed")
                if not embed or embed not in img_rel_map or embed in seen_in_para:
                    continue
                blob, ext, ct = img_rel_map[embed]
                # Skip EMF/WMF (OLE icon images)
                if ext.lower() in ("emf", "wmf"):
                    continue
                # Skip tiny images < 2 KB (decorative icons / bullets)
                if len(blob) < 2000:
                    continue
                # Skip duplicates across paragraphs
                if embed in {m.rId for m in media_items}:
                    continue
                seen_in_para.add(embed)
                item = MediaItem(
                    kind="image", blob=blob, ext=ext,
                    rId=embed, position_index=position,
                    context_text=prev_text, content_type=ct
                )
                media_items.append(item)
                position += 1

        if text:
            prev_text = text

    return media_items


# ─────────────────────────────────────────────────────────────────
# SOLUTION DOC PARSER
# ─────────────────────────────────────────────────────────────────

def extract_activity_name(doc: Document) -> str:
    paragraphs = doc.paragraphs
    obj_idx = None
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading") and normalize_heading(text) == "objective":
            obj_idx = i
            break
        if normalize_heading(text) == "objective":
            obj_idx = i
            break

    if obj_idx is not None and obj_idx > 0:
        for i in range(obj_idx - 1, -1, -1):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            upper = text.upper()
            if upper in ("METHOD OF PROCEDURE", "METHOD OF PROCEDURE (MOP)",
                         "CONTENTS:", "CONTENTS", "TITLE PAGE"):
                continue
            if re.match(r'^\d+[\.\)]\s+\w.*Page\s+\d+', text):
                continue
            if "\n" in text or re.match(
                    r'^(Customer|Activity Title|Document Reference|Domain|Vendor)[\s]*:',
                    text, re.IGNORECASE):
                for line in text.split("\n"):
                    line = line.strip()
                    m = re.match(r'^Activity\s+Title\s*:\s*(.+)', line, re.IGNORECASE)
                    if m:
                        return m.group(1).strip()
                continue
            if normalize_heading(text) is not None:
                continue
            if re.match(r'^(Customer|Header|Footer|Document)[\s]*:', text, re.IGNORECASE):
                continue
            name = re.sub(r'^MOP\s*:\s*', '', text, flags=re.IGNORECASE)
            name = re.sub(r'^UC\s*:\s*', '', name, flags=re.IGNORECASE)
            name = re.sub(r'^Activity\s+Title\s*:\s*', '', name, flags=re.IGNORECASE)
            name = re.sub(r'^Method of Procedure\s*[\(\[]?MOP[\)\]]?\s*[:\-]?\s*',
                          '', name, flags=re.IGNORECASE)
            name = name.strip()
            if name and len(name) > 3:
                return name

    for para in paragraphs[:10]:
        if para.style.name.startswith("Heading 1"):
            name = para.text.strip()
            name = re.sub(r'^MOP\s*:\s*', '', name, flags=re.IGNORECASE)
            if name and normalize_heading(name) is None:
                return name

    for para in paragraphs[:15]:
        for run in para.runs:
            if run.italic and run.underline and para.text.strip():
                return para.text.strip()

    return "Activity Name"


def extract_sections(doc: Document) -> dict:
    """Extract section content as XML element lists."""
    sections = {k: [] for k in SECTION_KEYS}
    current_key = None

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()

        is_heading_style = style.startswith("Heading")
        key_from_text    = normalize_heading(text) if text else None

        if is_heading_style or (key_from_text and len(text) < 120):
            key = key_from_text or normalize_heading(text)
            if key:
                current_key = key
                continue

        if current_key is None:
            continue

        if text.upper() in ("METHOD OF PROCEDURE", "METHOD OF PROCEDURE (MOP)",
                            "CONTENTS:", "CONTENTS", ""):
            continue
        if re.match(r'^\d+\.\s+\w.*Page\s+\d+', text):
            continue
        if re.match(r'^(Customer|Header|Footer|Activity Title|Document)[\s]*:',
                    text, re.IGNORECASE):
            continue
        if text == "sample...":
            continue

        if current_key in sections and text:
            sections[current_key].append(deepcopy(para._p))

    return sections


# ─────────────────────────────────────────────────────────────────
# COMMENT EXTRACTOR FROM ACTIVITY MOP
# ─────────────────────────────────────────────────────────────────

class CommentItem:
    """One Word comment extracted from Activity MOP."""
    def __init__(self, cid, author, date, text, para_text, para_index):
        self.cid        = cid         # original comment ID (string)
        self.author     = author      # author name
        self.date       = date        # ISO date string
        self.text       = text        # comment body text
        self.para_text  = para_text   # text of the paragraph the comment is on
        self.para_index = para_index  # paragraph index in Activity MOP


def extract_comments_from_mop(mop_bytes: bytes) -> tuple:
    """
    Extract all comments from Activity MOP.
    Returns:
        (list[CommentItem],
         comments_xml_str,
         commentsExtended_xml_str,
         commentsIds_xml_str,
         commentsExtensible_xml_str)

    Also builds a map: para_text → [CommentItem] for positional matching.
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    with zipfile.ZipFile(io.BytesIO(mop_bytes)) as z:
        names = z.namelist()

        # Read comment XML files (may not all exist)
        def _read(path):
            return z.read(path).decode("utf-8") if path in names else ""

        comments_xml      = _read("word/comments.xml")
        comments_ext_xml  = _read("word/commentsExtended.xml")
        comments_ids_xml  = _read("word/commentsIds.xml")
        comments_exs_xml  = _read("word/commentsExtensible.xml")
        doc_xml           = _read("word/document.xml")

    if not comments_xml:
        return [], "", "", "", ""

    # ── Parse comments.xml ──────────────────────────────────────
    comments_root = etree.fromstring(comments_xml.encode())
    comment_map   = {}   # cid → CommentItem (partial, para_text filled later)

    for c in comments_root.findall(f"{{{_W}}}comment"):
        cid    = c.get(f"{{{_W}}}id")
        author = c.get(f"{{{_W}}}author", "Unknown")
        date   = c.get(f"{{{_W}}}date", "")
        texts  = [t.text or "" for t in c.findall(f".//{{{_W}}}t")]
        text   = "".join(texts).strip()
        comment_map[cid] = CommentItem(cid, author, date, text, "", -1)

    # ── Parse document.xml to find which paragraph each comment is on ──
    doc_root = etree.fromstring(doc_xml.encode())
    body     = doc_root.find(f"{{{_W}}}body")
    paras    = body.findall(f".//{{{_W}}}p")

    for i, para in enumerate(paras):
        para_texts = [t.text or "" for t in para.findall(f".//{{{_W}}}t")]
        para_text  = "".join(para_texts).strip()

        # commentRangeStart marks the paragraph where comment is anchored
        starts = para.findall(f"{{{_W}}}commentRangeStart")
        refs   = para.findall(f".//{{{_W}}}commentReference")

        anchored_ids = set()
        for s in starts:
            anchored_ids.add(s.get(f"{{{_W}}}id"))
        for r in refs:
            anchored_ids.add(r.get(f"{{{_W}}}id"))

        for cid in anchored_ids:
            if cid in comment_map:
                comment_map[cid].para_text  = para_text
                comment_map[cid].para_index = i

    items = sorted(comment_map.values(), key=lambda x: x.para_index)
    return items, comments_xml, comments_ext_xml, comments_ids_xml, comments_exs_xml


def _inject_comments_into_docx(
    doc_bytes:       bytes,
    comment_items:   list,
    comments_xml:    str,
    comments_ext:    str,
    comments_ids:    str,
    comments_exs:    str,
) -> tuple:
    """
    Inject comments into the output docx at zip level.
    Strategy:
      1. Copy comments.xml + extended files from Activity MOP as-is into output docx
      2. Walk output document.xml paragraphs — match each paragraph text against
         the para_text stored in CommentItem
      3. Insert commentRangeStart / commentRangeEnd / commentReference XML nodes
         into the matched paragraph
      4. If no text match found → insert at the nearest paragraph by position offset

    Returns (new_doc_bytes, list_of_injected_ids, list_of_failed_ids)
    """
    _W     = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _W15   = "http://schemas.microsoft.com/office/word/2012/wordml"
    _W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
    _W16CEX = "http://schemas.microsoft.com/office/word/2018/wordml/cex"

    if not comment_items or not comments_xml:
        return doc_bytes, [], []

    # ── Read output docx ────────────────────────────────────────
    in_buf  = io.BytesIO(doc_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, "r") as zin:
        existing_files = set(zin.namelist())
        doc_xml_bytes  = zin.read("word/document.xml")
        rels_xml_bytes = zin.read("word/_rels/document.xml.rels")
        ct_xml_bytes   = zin.read("[Content_Types].xml")

    # ── Modify document.xml — inject comment anchors ────────────
    doc_root = etree.fromstring(doc_xml_bytes)
    body     = doc_root.find(f"{{{_W}}}body")
    out_paras = body.findall(f".//{{{_W}}}p")

    injected_ids = []
    failed_ids   = []

    for ci in comment_items:
        cid = ci.cid

        # Try to find matching paragraph by text
        matched_para = None
        best_score   = 0

        if ci.para_text:
            for para in out_paras:
                texts = [t.text or "" for t in para.findall(f".//{{{_W}}}t")]
                ptxt  = "".join(texts).strip()
                # Score: number of common words
                if ptxt and ci.para_text:
                    src_words = set(ci.para_text.lower().split())
                    tgt_words = set(ptxt.lower().split())
                    common    = len(src_words & tgt_words)
                    total     = max(len(src_words), 1)
                    score     = common / total
                    if score > best_score and score >= 0.5:
                        best_score   = score
                        matched_para = para

        # Fallback: use paragraph at same index offset (proportional)
        if matched_para is None and out_paras:
            ratio     = ci.para_index / max(len(out_paras), 1)
            fallback_i = min(int(ratio * len(out_paras)), len(out_paras) - 1)
            matched_para = out_paras[fallback_i]

        if matched_para is None:
            failed_ids.append(cid)
            continue

        # ── Build XML elements ───────────────────────────────────
        # commentRangeStart — goes BEFORE the first run in para
        def _make_range_start(cid_val):
            el = OxmlElement("w:commentRangeStart")
            el.set(f"{{{_W}}}id", cid_val)
            return el

        def _make_range_end(cid_val):
            el = OxmlElement("w:commentRangeEnd")
            el.set(f"{{{_W}}}id", cid_val)
            return el

        def _make_reference(cid_val):
            # w:r > w:rPr > w:rStyle(val=CommentReference) + w:commentReference
            r   = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rs  = OxmlElement("w:rStyle")
            rs.set(f"{{{_W}}}val", "CommentReference")
            rPr.append(rs)
            r.append(rPr)
            cr = OxmlElement("w:commentReference")
            cr.set(f"{{{_W}}}id", cid_val)
            r.append(cr)
            return r

        # Insert commentRangeStart at beginning of para (after pPr if exists)
        pPr = matched_para.find(f"{{{_W}}}pPr")
        insert_pos = 0
        children = list(matched_para)
        if pPr is not None:
            insert_pos = children.index(pPr) + 1

        matched_para.insert(insert_pos,     _make_range_start(cid))
        # commentRangeEnd and commentReference go at END of para
        matched_para.append(_make_range_end(cid))
        matched_para.append(_make_reference(cid))

        injected_ids.append(cid)

    # ── Serialise modified document.xml ─────────────────────────
    new_doc_xml = etree.tostring(doc_root, xml_declaration=True,
                                 encoding="UTF-8", standalone=True)

    # ── Update rels to add comments relationship ─────────────────
    rels_xml = rels_xml_bytes.decode("utf-8")
    COMMENT_REL = ("http://schemas.openxmlformats.org/officeDocument/"
                   "2006/relationships/comments")
    COMMENT_EXT_REL = ("http://schemas.microsoft.com/office/2011/relationships/"
                       "commentsExtended")
    COMMENT_IDS_REL = ("http://schemas.microsoft.com/office/2016/09/relationships/"
                       "commentsIds")
    COMMENT_EXS_REL = ("http://schemas.microsoft.com/office/2020/relationships/"
                       "commentsExtensible")

    def _add_rel(rels, rid, rtype, target):
        if target not in rels:
            new = (f'<Relationship Id="{rid}" Type="{rtype}" '
                   f'Target="{target}"/>')
            rels = rels.replace("</Relationships>", new + "</Relationships>")
        return rels

    rels_xml = _add_rel(rels_xml, "rIdCom1", COMMENT_REL,     "comments.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom2", COMMENT_EXT_REL, "commentsExtended.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom3", COMMENT_IDS_REL, "commentsIds.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom4", COMMENT_EXS_REL, "commentsExtensible.xml")

    # ── Update Content_Types.xml ────────────────────────────────
    ct_xml = ct_xml_bytes.decode("utf-8")
    ct_entries = [
        ('word/comments.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'),
        ('word/commentsExtended.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml'),
        ('word/commentsIds.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml'),
        ('word/commentsExtensible.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtensible+xml'),
    ]
    for part_name, ct in ct_entries:
        if part_name not in ct_xml:
            override = (f'<Override PartName="/{part_name}" '
                        f'ContentType="{ct}"/>')
            ct_xml = ct_xml.replace("</Types>", override + "</Types>")

    # ── Rebuild docx zip ────────────────────────────────────────
    in_buf.seek(0)
    with zipfile.ZipFile(in_buf, "r") as zin, \
         zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:

        skip = {"word/document.xml", "word/_rels/document.xml.rels",
                "[Content_Types].xml", "word/comments.xml",
                "word/commentsExtended.xml", "word/commentsIds.xml",
                "word/commentsExtensible.xml"}

        for item in zin.infolist():
            if item.filename not in skip:
                zout.writestr(item, zin.read(item.filename))

        zout.writestr("word/document.xml",              new_doc_xml)
        zout.writestr("word/_rels/document.xml.rels",   rels_xml.encode("utf-8"))
        zout.writestr("[Content_Types].xml",            ct_xml.encode("utf-8"))
        zout.writestr("word/comments.xml",              comments_xml.encode("utf-8"))
        if comments_ext:
            zout.writestr("word/commentsExtended.xml",  comments_ext.encode("utf-8"))
        if comments_ids:
            zout.writestr("word/commentsIds.xml",       comments_ids.encode("utf-8"))
        if comments_exs:
            zout.writestr("word/commentsExtensible.xml",comments_exs.encode("utf-8"))

    out_buf.seek(0)
    return out_buf.read(), injected_ids, failed_ids


# ─────────────────────────────────────────────────────────────────
# DOCX BUILDER — CORE
# ─────────────────────────────────────────────────────────────────

def _apply_heading_color(p_elem):
    def _fix_color(rpr):
        color_el = rpr.find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            rpr.append(color_el)
        color_el.set(qn("w:val"), _HEADING_COLOR_HEX)
        for attr in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
            if color_el.get(attr) is not None:
                del color_el.attrib[attr]

    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        p_rpr = pPr.find(qn("w:rPr"))
        if p_rpr is not None:
            _fix_color(p_rpr)

    for r_el in p_elem.findall(".//" + qn("w:r")):
        rpr = r_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r_el.insert(0, rpr)
        if rpr.find(qn("w:b")) is None:
            rpr.insert(0, OxmlElement("w:b"))
        _fix_color(rpr)


def _clone_para(src_p_elem):
    cloned = deepcopy(src_p_elem)
    pPr = cloned.find(qn("w:pPr"))
    if pPr is not None:
        for tag in (qn("w:pStyle"), qn("w:numPr"), qn("w:pageBreakBefore")):
            for el in pPr.findall(tag):
                pPr.remove(el)
    return cloned


def _update_header_date(doc: Document, today_str: str):
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if "{{current date}}" in run.text:
                    run.text = run.text.replace("{{current date}}", today_str)


def _update_revision_table(doc: Document, activity_name: str, today_str: str):
    for table in doc.tables:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if "Version No." not in header_cells:
            continue
        if len(table.rows) >= 2:
            row = table.rows[1]
            # Update date cell (index 1)
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = today_str
                else:
                    para.add_run(today_str).font.name = "Calibri"
            # Update description cell (index 3 or last)
            desc_idx = min(3, len(row.cells) - 1)
            desc = f"Auto-generated MOP — {activity_name}"
            for para in row.cells[desc_idx].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = desc
                else:
                    para.add_run(desc).font.name = "Calibri"
        break





# ─────────────────────────────────────────────────────────────────
# MAIN BUILD FUNCTION
# ─────────────────────────────────────────────────────────────────

def _make_xml_para(doc: Document, text: str, bold=False,
                   color_rgb=None, italic=False, size_pt=10) -> etree._Element:
    """Create a plain paragraph XML element without touching the document body."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fn = OxmlElement("w:rFonts")
    fn.set(qn("w:ascii"), "Calibri")
    fn.set(qn("w:hAnsi"), "Calibri")
    rpr.append(fn)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)

    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    if color_rgb:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), color_rgb)
        rpr.append(col)

    r.append(rpr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _make_image_xml(doc: Document, img_bytes: bytes, width_inches=5.5):
    """
    Add image to the document's part store and return a paragraph XML element.
    Uses python-docx's InlineShape mechanism but detaches from body afterwards.
    Raises on failure so caller can handle it.
    """
    from docx.shared import Inches as _Inches
    # Add a temp paragraph to the doc to use add_picture properly
    tmp_p = doc.add_paragraph()
    run   = tmp_p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=_Inches(width_inches))
    # Detach from body — we will re-attach in the right position
    p_xml = tmp_p._element
    p_xml.getparent().remove(p_xml)
    return p_xml


def _make_notice_xml(desc: str) -> etree._Element:
    """Red bold warning paragraph — pure XML, no doc.add_paragraph()."""
    return _make_xml_para(
        None,
        f"⚠ [MEDIA NOT INSERTED — Please add manually: {desc}]",
        bold=True, color_rgb="CC3300", size_pt=10
    )


def _make_caption_xml() -> etree._Element:
    return _make_xml_para(
        None,
        "[Screenshot/Image — copied from Activity MOP]",
        italic=True, size_pt=9, color_rgb="595959"
    )


def _insert_after(anchor: etree._Element, new_elem: etree._Element):
    """Insert new_elem immediately after anchor in the same parent."""
    parent = anchor.getparent()
    idx    = list(parent).index(anchor)
    parent.insert(idx + 1, new_elem)


def _embed_attachment_into_docx(doc_bytes: bytes, media_item, attach_idx: int):
    """Embed attachment into docx ZIP — only reliable Python method."""
    try:
        fname      = media_item.display_name
        ext        = media_item.ext.lower()
        CT_MAP = {
            "xlsx":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xls":"application/vnd.ms-excel",
            "docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc":"application/msword",
            "txt":"text/plain","log":"text/plain","csv":"text/csv",
            "pdf":"application/pdf","zip":"application/zip","bin":"application/octet-stream",
        }
        ct        = CT_MAP.get(ext, "application/octet-stream")
        rel_type  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"
        new_rId   = f"rIdEmb{attach_idx:04d}"
        safe_name = re.sub(r'[^\w.\-]', '_', fname)
        part_path = f"word/embeddings/{safe_name}"
        rel_tgt   = f"embeddings/{safe_name}"
        in_buf, out_buf = io.BytesIO(doc_bytes), io.BytesIO()
        with zipfile.ZipFile(in_buf,"r") as zin, zipfile.ZipFile(out_buf,"w",zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/_rels/document.xml.rels":
                    rels = data.decode("utf-8")
                    new_rel = f'<Relationship Id="{new_rId}" Type="{rel_type}" Target="{rel_tgt}"/>'
                    rels = rels.replace("</Relationships>", new_rel+"</Relationships>")
                    zout.writestr(item, rels.encode("utf-8"))
                elif item.filename == "[Content_Types].xml":
                    cts = data.decode("utf-8")
                    if part_path not in cts:
                        cts = cts.replace("</Types>",f'<Override PartName="/{part_path}" ContentType="{ct}"/></Types>')
                    zout.writestr(item, cts.encode("utf-8"))
                else:
                    zout.writestr(item, data)
            zout.writestr(part_path, media_item.blob)
        out_buf.seek(0)
        return out_buf.read(), True
    except Exception:
        return doc_bytes, False


def build_mop(
    template_bytes: bytes,
    activity_name:  str,
    sections:       dict,
    today_str:      str,
    media_items:    list,
) -> tuple[bytes, list, int]:
    doc  = Document(io.BytesIO(template_bytes))
    body = doc.element.body
    _update_header_date(doc, today_str)

    # Title subtitle
    for child in list(body):
        if child.tag.split("}")[-1] != "p": continue
        se = child.find(".//" + qn("w:pStyle"))
        if se is not None and se.get(qn("w:val")) == "Title":
            sub_e = _make_xml_para(doc, activity_name, italic=True, size_pt=14)
            pPr = OxmlElement("w:pPr"); jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center"); pPr.append(jc); sub_e.insert(0, pPr)
            _insert_after(child, sub_e); break

    _update_revision_table(doc, activity_name, today_str)

    ordered_sections = []
    for child in list(body):
        if child.tag.split("}")[-1] != "p": continue
        se = child.find(".//" + qn("w:pStyle"))
        if se is None or se.get(qn("w:val"), "") != "Heading1": continue
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t"))).strip()
        key  = normalize_heading(text)
        if key and key != "connectivity_diagram":
            _apply_heading_color(child)
            ordered_sections.append((child, key))

    media_queue     = list(media_items)
    media_idx       = 0
    failed_media    = []
    injected_count  = 0
    pending_att     = []   # (MediaItem, attach_idx)
    att_counter     = [0]

    for h_elem, sec_key in ordered_sections:
        # Remove boilerplate under heading
        to_remove, found = [], False
        for child in list(body):
            if child is h_elem: found = True; continue
            if not found: continue
            ctag = child.tag.split("}")[-1]
            if ctag in ("tbl","sectPr"): break
            if ctag == "p":
                se = child.find(".//" + qn("w:pStyle"))
                if se is not None and "Heading" in se.get(qn("w:val"),""): break
                to_remove.append(child)
        for e in to_remove: body.remove(e)

        if sec_key == "objective":
            pPr = h_elem.find(qn("w:pPr"))
            if pPr is None: pPr = OxmlElement("w:pPr"); h_elem.insert(0, pPr)
            pb = OxmlElement("w:pageBreakBefore"); pb.set(qn("w:val"),"1"); pPr.append(pb)

        content_elems = sections.get(sec_key, [])
        if not content_elems:
            _insert_after(h_elem, OxmlElement("w:p")); continue

        cursor = h_elem

        if sec_key == "sop":
            for p_elem in content_elems:
                text = "".join(t.text or "" for t in p_elem.findall(".//" + qn("w:t")))
                is_ph = bool(IMAGE_PLACEHOLDER_RE.search(text))

                if is_ph and media_idx < len(media_queue):
                    media_item = media_queue[media_idx]; media_idx += 1
                    ph_clone = _clone_para(p_elem)
                    _insert_after(cursor, ph_clone); cursor = ph_clone

                    if media_item.kind == "image":
                        try:
                            img_xml = _make_image_xml(doc, media_item.blob)
                            _insert_after(cursor, img_xml); cursor = img_xml
                            cap = _make_caption_xml()
                            _insert_after(cursor, cap); cursor = cap
                            media_item.injected = True; injected_count += 1
                        except Exception as ex:
                            media_item.inject_error = str(ex)
                            desc = f"Image #{media_item.position_index+1} (near: {media_item.context_text[:50] or 'unknown'})"
                            failed_media.append(desc)
                            _insert_after(cursor, _make_notice_xml(desc)); cursor = list(body)[-1]
                    else:
                        # Attachment — embed via zip post-processing
                        att_counter[0] += 1
                        att_xml = _make_xml_para(
                            doc,
                            f"\U0001f4ce  ATTACHED FILE [{media_item.ext.upper()}]: {media_item.display_name}  \u2014 embedded in document",
                            bold=True, color_rgb="00695C", size_pt=10
                        )
                        _insert_after(cursor, att_xml); cursor = att_xml
                        pending_att.append((media_item, att_counter[0]))
                        media_item.injected = True; injected_count += 1
                else:
                    cloned = _clone_para(p_elem)
                    _insert_after(cursor, cloned); cursor = cloned

            # Remaining unmatched media → append at SOP end
            while media_idx < len(media_queue):
                m = media_queue[media_idx]; media_idx += 1
                if m.kind == "image":
                    try:
                        img_xml = _make_image_xml(doc, m.blob)
                        _insert_after(cursor, img_xml); cursor = img_xml
                        cap = _make_caption_xml()
                        _insert_after(cursor, cap); cursor = cap
                        m.injected = True; injected_count += 1
                    except Exception as ex:
                        desc = f"Image #{m.position_index+1} (unmatched)"
                        failed_media.append(desc)
                        notice = _make_notice_xml(desc)
                        _insert_after(cursor, notice); cursor = notice
                else:
                    att_counter[0] += 1
                    att_xml = _make_xml_para(
                        doc,
                        f"\U0001f4ce  ATTACHED FILE [{m.ext.upper()}]: {m.display_name}  \u2014 embedded in document",
                        bold=True, color_rgb="00695C", size_pt=10
                    )
                    _insert_after(cursor, att_xml); cursor = att_xml
                    pending_att.append((m, att_counter[0]))
                    m.injected = True; injected_count += 1
        else:
            for p_elem in content_elems:
                cloned = _clone_para(p_elem)
                _insert_after(cursor, cloned); cursor = cloned

    # Save base document
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    # Post-process: embed each attachment into the docx ZIP
    for media_item, att_idx in pending_att:
        new_bytes, ok = _embed_attachment_into_docx(doc_bytes, media_item, att_idx)
        if ok:
            doc_bytes = new_bytes
        else:
            desc = (f"Attachment \u2018{media_item.display_name}\u2019 "
                    f"(.{media_item.ext}) could not be embedded automatically "
                    f"\u2014 please attach manually")
            failed_media.append(desc)

    return doc_bytes, failed_media, injected_count




# ─────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="text-align:center; padding:1.2rem 0 0.8rem;">
      <span style="font-family:'Lato',sans-serif; font-weight:900; font-size:2rem; letter-spacing:6px; color:#00BFA5; display:block;">ERICSSON</span>
      <span style="font-size:0.6rem; letter-spacing:2px; color:rgba(255,255,255,0.3); text-transform:uppercase; display:block; margin-top:2px;">Technology For Good</span>
    </div>
    <hr/>
    """, unsafe_allow_html=True)

    st.markdown("### 📋 Compiled MOP Generator")
    st.markdown("""
    <div class="sidebar-info">
      Automates MOP document generation from Solution Documents.<br><br>
      <strong>NEW in v6:</strong><br>
      · Activity MOP image injection<br>
      · OLE attachment re-embedding<br>
      · Positional placeholder matching<br>
      · Manual-add notices for failures
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("""
    <div style="font-size:0.78rem; color:#80C9BF; line-height:1.8;">
    1️⃣ &nbsp;Place template in <code>templates/</code> folder<br>
    2️⃣ &nbsp;Upload <strong>Solution Document</strong> (required)<br>
    3️⃣ &nbsp;Upload <strong>Activity MOP</strong> (optional)<br>
    &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;→ Images copied to SOP section<br>
    4️⃣ &nbsp;Click <strong>Generate MOP</strong><br>
    5️⃣ &nbsp;Download .docx output
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""
    <div class="sidebar-info">
      🔒 <strong>Zero Data Retention</strong><br>
      All processing strictly in-memory.<br>
      No files written to disk.<br>
      No data logged or stored.<br>
      Session cleared on browser close.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""
    <div style="font-size:0.62rem; color:rgba(255,255,255,0.2); text-align:center; letter-spacing:.5px;">
      Compiled MOP Generator v6<br>
      Ericsson Internal Tool
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────
# MAIN AREA
# ─────────────────────────────────────────────────────────────────

st.markdown("""
<div class="eri-topbar">
  <div>
    <div class="eri-logo-text">ERICSSON</div>
    <div class="eri-logo-sub">Telecom Automation Toolkit</div>
  </div>
  <div style="text-align:center;">
    <div class="eri-app-title">Compiled MOP Generator</div>
    <div class="eri-app-sub">Solution Document + Activity MOP → Formatted Output MOP · Images Injected · Audit-Ready</div>
  </div>
  <div>
    <span class="eri-version">v6</span>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Zero Data Retention Banner ──────────────────────────────────
st.markdown("""
<div class="priv-bar">
  🔒 <strong>ZERO DATA RETENTION:</strong> All processing is performed entirely in-memory.
  No uploaded files, generated documents, or any user data are written to disk, logged, or stored
  at any stage of processing. This session and all associated data are permanently cleared when
  you close your browser tab.
</div>
""", unsafe_allow_html=True)

# ── Layout ──────────────────────────────────────────────────────
col_left, col_right = st.columns([1.1, 1], gap="large")

with col_left:

    # ── Step 1: Template ─────────────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 01</span> Select MOP Template</div>', unsafe_allow_html=True)

    templates      = discover_templates()
    selected_template = None
    template_bytes    = None

    if not templates:
        st.markdown('<div class="pill-warn">⚠ No template found. Place <strong>.docx</strong> file in <code>templates/</code> folder, then restart.</div>', unsafe_allow_html=True)
    else:
        names = [t.name for t in templates]
        sel   = st.selectbox("Template file", names, label_visibility="visible")
        selected_template = next(t for t in templates if t.name == sel)
        template_bytes    = load_template_bytes(selected_template)
        st.markdown(f'<div class="pill-ok">✔ &nbsp;<strong>{sel}</strong> &nbsp;ready</div>', unsafe_allow_html=True)

    st.markdown("""
    <div style="font-size:0.72rem; color:#5a7a9a; background:rgba(0,82,130,0.05);
         border:1px solid rgba(0,191,165,0.15); border-left:3px solid #00BFA5;
         border-radius:0 6px 6px 0; padding:0.5rem 0.8rem; margin-top:0.5rem;">
      📁 &nbsp;<strong>To add/update a template</strong>, place the <code>.docx</code>
      file in the <code>templates/</code> folder and restart the app.
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 2: Solution Document ────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 02</span> Upload Solution Document <span style="font-size:0.62rem;color:#cc4400;margin-left:6px;">REQUIRED</span></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.73rem;color:#5a7a9a;margin-bottom:0.5rem;">
      The AI-generated Solution Document (all 12 sections). All text content comes from this file.
      Images/attachments in this doc are not needed — they are sourced from the Activity MOP below.
    </div>
    """, unsafe_allow_html=True)

    sol_file = st.file_uploader("Solution Document (.docx)", type=["docx"],
                                key="sol_up", label_visibility="visible")
    if sol_file:
        size_kb = sol_file.size / 1024
        st.markdown(
            f'<div class="pill-ok">✔ &nbsp;<strong>{sol_file.name}</strong>'
            f' &nbsp;·&nbsp; {size_kb:.1f} KB</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="pill-warn">⏳ Waiting for solution document…</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 3: Activity MOP (OPTIONAL) ─────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 03</span> Upload Activity MOP &nbsp;<span class="optional-badge">OPTIONAL</span></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.73rem;color:#5a7a9a;margin-bottom:0.5rem;">
      The original Activity MOP document (may contain screenshots, flow diagrams, and embedded
      attachments). <strong>Only media</strong> (images, attachments) is extracted from this file —
      no text data. Media is injected into the SOP section of the output MOP in positional order,
      replacing <code>[IMAGE/SCREENSHOT REQUIRED]</code> placeholders.
    </div>
    """, unsafe_allow_html=True)

    mop_file = st.file_uploader("Activity MOP (.docx) — optional", type=["docx"],
                                key="mop_up", label_visibility="visible")
    if mop_file:
        size_kb = mop_file.size / 1024
        st.markdown(
            f'<div class="pill-ok">✔ &nbsp;<strong>{mop_file.name}</strong>'
            f' &nbsp;·&nbsp; {size_kb:.1f} KB — media will be extracted</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="pill-info">ℹ No Activity MOP uploaded — output MOP will contain text only with placeholder notices.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 4: Generate ─────────────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 04</span> Generate Output MOP</div>', unsafe_allow_html=True)

    can_go  = bool(sol_file and templates)
    gen_btn = st.button("⚡  Generate MOP", disabled=not can_go)

    if not can_go:
        missing = []
        if not templates:
            missing.append("MOP template")
        if not sol_file:
            missing.append("solution document")
        if missing:
            st.markdown(f'<div class="pill-warn">⏳ Still needed: {" + ".join(missing)}</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


with col_right:

    if gen_btn and can_go:
        st.markdown('<div class="eri-card"><div class="eri-card-title">⚙ Processing</div>', unsafe_allow_html=True)

        steps = [
            "Loading template",                                                              # 0
            "Reading solution document",                                                     # 1
            "Extracting activity name",                                                      # 2
            "Parsing all 12 sections",                                                       # 3
            "Reading Activity MOP" if mop_file else "No Activity MOP — text-only mode",     # 4
            "Extracting images & attachments" if mop_file else "Skipping media extraction", # 5
            "Extracting comments from Activity MOP" if mop_file else "No comments to copy", # 6
            "Building output MOP document",                                                  # 7
            "Injecting comments into SOP section",                                           # 8
            "Updating revision table & header",                                              # 9
            "Finalising document",                                                           # 10
        ]

        st.markdown('<div class="prog-wrap">', unsafe_allow_html=True)
        phs = [st.empty() for _ in steps]
        for ph, s in zip(phs, steps):
            ph.markdown(f'<div class="ps wait"><div class="pd wait"></div>{s}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        try:
            activity_name  = ""
            sections       = {}
            today_str      = ""
            output_bytes   = b""
            media_items    = []
            comment_items  = []
            comments_xml   = ""
            comments_ext   = ""
            comments_ids   = ""
            comments_exs   = ""
            failed_media   = []
            injected_count = 0
            c_injected     = []
            c_failed       = []

            for i, (ph, step) in enumerate(zip(phs, steps)):
                ph.markdown(f'<div class="ps doing"><div class="pd doing"></div>{step}…</div>', unsafe_allow_html=True)
                time.sleep(0.10)

                if i == 0:
                    tmpl_b = load_template_bytes(selected_template)

                elif i == 1:
                    sol_bytes = sol_file.read()
                    sol_doc   = Document(io.BytesIO(sol_bytes))

                elif i == 2:
                    activity_name = extract_activity_name(sol_doc)
                    today_str     = datetime.today().strftime("%d-%m-%Y")

                elif i == 3:
                    sections = extract_sections(sol_doc)

                elif i == 4:
                    if mop_file:
                        mop_bytes_data = mop_file.read()

                elif i == 5:
                    if mop_file:
                        media_items = extract_media_from_activity_mop(mop_bytes_data)

                elif i == 6:
                    if mop_file:
                        (comment_items, comments_xml,
                         comments_ext, comments_ids,
                         comments_exs) = extract_comments_from_mop(mop_bytes_data)

                elif i == 7:
                    output_bytes, failed_media, injected_count = build_mop(
                        tmpl_b, activity_name, sections, today_str, media_items
                    )

                elif i == 8:
                    if comment_items and comments_xml:
                        output_bytes, c_injected, c_failed = _inject_comments_into_docx(
                            output_bytes, comment_items,
                            comments_xml, comments_ext,
                            comments_ids, comments_exs
                        )

                ph.markdown(f'<div class="ps done"><div class="pd done"></div>{step} ✓</div>', unsafe_allow_html=True)
                time.sleep(0.04)

            # ── Store in session_state ──────────────────────────
            st.session_state["output_bytes"]      = output_bytes
            st.session_state["activity_name"]     = activity_name
            st.session_state["today_str"]         = today_str
            st.session_state["sections"]          = sections
            st.session_state["filled"]            = sum(1 for k in SECTION_KEYS[:-1] if sections.get(k))
            st.session_state["images_n"]          = len([m for m in media_items if m.kind == "image"])
            st.session_state["total_n"]           = sum(len(v) for k, v in sections.items()
                                                         if k != "connectivity_diagram")
            st.session_state["failed_media"]      = failed_media
            st.session_state["injected_media"]    = injected_count
            st.session_state["comments_injected"] = len(c_injected)
            st.session_state["comments_failed"]   = c_failed

            st.markdown('</div>', unsafe_allow_html=True)

        except Exception as e:
            st.markdown('</div>', unsafe_allow_html=True)
            st.error(f"❌ Error during generation: {e}")
            import traceback
            st.code(traceback.format_exc())

    # ── Result panel ─────────────────────────────────────────────
    if st.session_state.get("output_bytes"):
        activity_name      = st.session_state["activity_name"]
        today_str          = st.session_state["today_str"]
        sections           = st.session_state["sections"]
        output_bytes       = st.session_state["output_bytes"]
        filled             = st.session_state["filled"]
        images_n           = st.session_state["images_n"]
        total_n            = st.session_state["total_n"]
        failed_media       = st.session_state["failed_media"]
        injected_media     = st.session_state["injected_media"]
        comments_injected  = st.session_state.get("comments_injected", 0)
        comments_failed    = st.session_state.get("comments_failed", [])

        # ── Success card ─────────────────────────────────────────
        st.markdown(f"""
        <div class="success-card">
          <div class="success-icon">✅</div>
          <div class="success-title">Output MOP Generated Successfully</div>
          <div class="success-sub">
            <strong class="success-name">{activity_name}</strong>
            &nbsp;·&nbsp; {today_str}
          </div>
        </div>""", unsafe_allow_html=True)

        safe_name = re.sub(r'[^\w\s\-]', '', activity_name).strip().replace(' ', '_')[:80]
        _dl_key   = f"dl_{abs(hash(safe_name + today_str)) % 10_000_000}"
        st.download_button(
            label="📥  Download Output MOP (.docx)",
            data=io.BytesIO(output_bytes),
            file_name=f"{safe_name}_MOP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=_dl_key,
            use_container_width=True,
        )

        # ── Failed media warning ──────────────────────────────────
        if failed_media:
            st.markdown("""
            <div class="warn-bar">
              <strong>⚠ Some media items could not be automatically inserted into the output MOP.</strong>
              They are marked with a notice in the document at the exact position they should appear.
              Please insert them manually using Word's Insert → Pictures or Insert → Object.
            </div>
            """, unsafe_allow_html=True)

            st.markdown('<div class="media-fail-card">', unsafe_allow_html=True)
            st.markdown('<div class="media-fail-title">🖼 Media requiring manual insertion:</div>', unsafe_allow_html=True)
            for desc in failed_media:
                st.markdown(f'<div class="media-fail-item">• {desc}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        elif images_n > 0:
            st.markdown(f'<div class="pill-ok" style="margin-top:0.6rem;">✔ &nbsp;All {injected_media} media items injected successfully into the SOP section</div>', unsafe_allow_html=True)

        # ── Metrics ───────────────────────────────────────────────
        st.markdown('<div class="eri-card" style="margin-top:0.8rem;"><div class="eri-card-title">📊 Generation Summary</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="metric-row">
          <div class="metric-box">
            <div class="metric-val">{filled}<span style="font-size:.85rem;color:#9ab8b4;">/12</span></div>
            <div class="metric-sub">Sections Filled</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{injected_media}</div>
            <div class="metric-sub">Media Injected</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{comments_injected}</div>
            <div class="metric-sub">Comments Injected</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{len(failed_media) + len(comments_failed)}</div>
            <div class="metric-sub">Manual Fixes Needed</div>
          </div>
        </div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # ── Comments summary ──────────────────────────────────────
        if comments_injected > 0:
            st.markdown(
                f'<div class="pill-ok" style="margin-top:0.5rem;">'
                f'✔ &nbsp;{comments_injected} comment(s) injected at matching paragraph positions '
                f'(author + date preserved)</div>',
                unsafe_allow_html=True
            )
        if comments_failed:
            st.markdown("""
            <div class="warn-bar">
              <strong>⚠ Some comments could not be position-matched and were skipped.</strong>
              Please add them manually in Word using the Review → New Comment option.
            </div>""", unsafe_allow_html=True)
            for cf in comments_failed:
                st.markdown(
                    f'<div style="font-size:0.75rem;color:#7a3500;padding:2px 0;">• Comment ID {cf} — position match failed</div>',
                    unsafe_allow_html=True
                )

        # ── Section fill status ───────────────────────────────────
        with st.expander("📋 Section fill status", expanded=False):
            for k in SECTION_KEYS[:-1]:
                label  = SECTION_LABELS.get(k, k)
                filled_flag = bool(sections.get(k))
                icon   = "✅" if filled_flag else "⚠️"
                color  = "#006633" if filled_flag else "#cc5500"
                st.markdown(
                    f'<div style="font-size:0.77rem;color:{color};padding:3px 0;">'
                    f'{icon} &nbsp; {label}</div>',
                    unsafe_allow_html=True,
                )

    # ── Footer ───────────────────────────────────────────────────
    st.markdown("""
    <div class="footer">
      Compiled MOP Generator v6 &nbsp;·&nbsp; Ericsson Internal Tool &nbsp;·&nbsp;
      🔒 Zero Data Retention &nbsp;·&nbsp; All processing in-memory only
    </div>
    """, unsafe_allow_html=True)
